# use ELN-procedure-venv
import pandas as pd
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document

# store a csv as datadfame TODO accept user input
# file_name = 'procedure_00'
file_name = input('file name: ')
procedure = pd.read_csv(f'{file_name}.csv')
print(procedure)

# extracting only reagents from the df where time is NaN or 'set-up'
procedure_reagents = procedure[(procedure.type=='reagent') & ((procedure.time.isna()) | (procedure.time == 'set-up'))]

# slicing the df for relevant columns
procedure_reagents = procedure_reagents[['description','details']]

# splitting the details into an value and a unit and renaming apropriatelly
procedure_reagents[['value', 'unit']] = procedure_reagents['details'].str.split(' ', expand=True)
# procedure_reagents.drop(['details'], axis = 1, inplace = True)
procedure_reagents.rename(columns={'description': 'nickname', 'details': 'amount'}, inplace=True)
print(procedure_reagents)

# adding a sub_phrase column
procedure_reagents['sub_phrase'] = procedure_reagents.apply(lambda row: f'{str(row["amount"])} of {row["nickname"]}',axis = 1)
print(procedure_reagents)

# ini w/ a dict. TODO load from a csv file to a Dataframe, pivot table and finally take a single row as a regaent {'nickname': 'NaH', 'amount_value': 100, 'amount_unit': 'mg' }

# the beginning of the sentence
charge_phrase = 'Reaction vessel was charged with '

# reagents unpacking
print(procedure_reagents['sub_phrase'])

# reagents = procedure['reagents']
sub_phrases = procedure_reagents['sub_phrase'].tolist()
if len(sub_phrases) > 1:
    charge_phrase += ', '.join(sub_phrases[:-1]) 
    charge_phrase += f' and {sub_phrases[-1]}.' 
else:
    charge_phrase += f'{sub_phrases[-1]}.' 

print(charge_phrase)

# export to word document with docxtpl
# document_path = Path(__file__).parent / "procedure_00.docx"
doc = Document()
paragraph = doc.add_paragraph()
paragraph.add_run(charge_phrase)
doc.save(f'{file_name}.docx')
print(f'procedure {file_name} was saved to a word document')

